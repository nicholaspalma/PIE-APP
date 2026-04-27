import streamlit as st
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from io import BytesIO
import os
import base64

# --- 1. CONFIGURACIÓN INICIAL Y UI ---
st.set_page_config(page_title="PACI Pro: Edición Infantil", page_icon="🎨", layout="wide")

# Función para el fondo (Mejorada para manejar errores)
def apply_custom_ui(image_file):
    try:
        if os.path.exists(image_file):
            with open(image_file, "rb") as f:
                encoded_string = base64.b64encode(f.read()).decode()
            st.markdown(f"""
                <style>
                .stApp {{
                    background-image: url("data:image/png;base64,{encoded_string}");
                    background-size: cover;
                    background-attachment: fixed;
                }}
                .top-header {{
                    background-color: rgba(255, 255, 255, 0.9); border-radius: 15px; padding: 15px;
                    text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                    margin-bottom: 20px;
                }}
                .item-card {{
                    background-color: rgba(255, 255, 255, 0.95);
                    border-radius: 15px; padding: 25px; margin-bottom: 20px;
                    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                }}
                [data-testid="stSidebar"] {{ display: none; }}
                </style>
                """, unsafe_allow_html=True)
        else:
             # Si no hay imagen, un fondo azul muy suave
             st.markdown("""<style>.stApp { background-color: #f0f8ff; }</style>""", unsafe_allow_html=True)
    except Exception as e:
        st.warning(f"No se pudo cargar el fondo: {e}")

# Aplicar el fondo
apply_custom_ui("fondo.png")

st.markdown('<div class="top-header"><h1>🎨 PACI Pro: Nivel Gemini 3</h1></div>', unsafe_allow_html=True)

# --- 2. CONFIGURACIÓN DEL MENÚ ---
st.markdown('<div class="item-card">', unsafe_allow_html=True)
st.markdown("### 1. Configura la adecuación")
col1, col2, col3 = st.columns(3)
with col1: asignatura_sel = st.selectbox("Asignatura:", ["Lenguaje", "Matemáticas", "Historia", "Ciencias", "Inglés"])
with col2: curso_sel = st.selectbox("Curso:", ["1ro Básico", "2do Básico", "3ro Básico", "4to Básico", "5to Básico", "6to Básico", "7mo Básico", "8vo Básico"])
with col3: necesidad_sel = st.selectbox("Necesidad (PIE):", ["TEA", "TDAH", "Trastorno del Lenguaje"])
st.markdown('</div>', unsafe_allow_html=True)

# --- 3. LECTURA DEL DOCUMENTO ---
def leer_docx(archivo):
    doc = docx.Document(archivo)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# --- 4. CREACIÓN DEL DOCUMENTO ADAPTADO (MEJORADA) ---
def crear_docx_adaptado(texto_adaptado, titulo_documento):
    # Cargar plantilla o crear uno nuevo
    if os.path.exists("plantilla.docx"):
        doc = docx.Document("plantilla.docx")
        # Buscar y reemplazar {titulo}
        for p in doc.paragraphs:
            if "{titulo}" in p.text:
                p.text = p.text.replace("{titulo}", titulo_documento)
                for run in p.runs: 
                    run.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 0, 0) # Asegurar color negro
    else:
        doc = docx.Document()
        # Si no hay plantilla, agregar el título arriba
        t = doc.add_paragraph(titulo_documento)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.runs[0].bold = True
        t.runs[0].font.size = Pt(16)
        doc.add_paragraph("") # Espacio en blanco

    # Procesar el texto de la IA
    lineas = [l.strip() for l in texto_adaptado.split('\n') if l.strip()]
    letra_opcion = 0
    abc = ["a)", "b)", "c)", "d)", "e)"]
    dibujo_creado = False

    for linea in lineas:
        # 1. Ignorar el título si la IA lo repitió al principio
        if linea.upper() == titulo_documento.upper() or linea.replace("#", "").strip().upper() == titulo_documento.upper():
            continue

        # 2. Limpiar la línea de formatos indeseados de Markdown
        linea_limpia = linea.replace('*', '').replace('_', '').replace('#', '').strip()
        
        if not linea_limpia: continue # Saltar líneas que quedaron vacías tras la limpieza

        # 3. Detectar instrucciones de DIBUJO
        if ("DIBUJA" in linea_limpia.upper() or "DIBUJO" in linea_limpia.upper()) and not dibujo_creado:
            p = doc.add_paragraph(f"🎨 {linea_limpia}")
            p.runs[0].bold = True
            p.paragraph_format.space_after = Pt(12)
            
            # Crear un recuadro (tabla de 1x1) para el dibujo
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid' # Borde estándar
            cell = table.rows[0].cells[0]
            # Añadir 8 saltos de línea para hacer el recuadro grande
            for _ in range(8): cell.add_paragraph("")
            
            doc.add_paragraph("") # Espacio después del recuadro
            dibujo_creado = True
            continue

        # 4. Detectar instrucciones de ESCRITURA (Cuadrícula)
        if "ESCRIBE" in linea_limpia.upper() or "COMPLETA" in linea_limpia.upper():
            p = doc.add_paragraph(f"✏️ {linea_limpia}")
            p.runs[0].bold = True
            p.paragraph_format.space_after = Pt(12)
            
            # Crear una tabla simple para simular la cuadrícula
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            for cell in table.rows[0].cells:
                 cell.width = Inches(0.5) # Ancho uniforme
            doc.add_paragraph("")
            continue

        # 5. Formato de Alternativas (Convertir guiones en letras)
        if linea_limpia.startswith(('-', '•', '>')):
            # Si empieza con guion o viñeta, asignar letra
            texto_final = f"      {abc[letra_opcion % 5]} {linea_limpia[1:].strip()}"
            letra_opcion += 1
            p = doc.add_paragraph(texto_final)
            p.paragraph_format.space_after = Pt(6)
            continue
        else:
            # Es un párrafo normal o una nueva pregunta, reiniciar letras
            letra_opcion = 0
            p = doc.add_paragraph(linea_limpia)
            p.paragraph_format.space_after = Pt(12)
            
            # Destacar partes o secciones
            if "PARTE" in linea_limpia.upper() or "ÍTEM" in linea_limpia.upper():
                p.runs[0].bold = True
                p.paragraph_format.space_before = Pt(18)

    # Guardar en memoria
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- 5. LLAMADA A LA IA (PROMPT MEJORADO) ---
def llamar_ia(texto, curso, asignatura, necesidad, api_key, modelo_nombre):
    genai.configure(api_key=api_key)
    modelo = genai.GenerativeModel(modelo_nombre)
    
    prompt = f"""
    Eres una Educadora Diferencial experta. Tu tarea es adaptar esta evaluación de {asignatura} para un estudiante de {curso} con {necesidad}.

    REGLAS ESTRICTAS DE RESPUESTA:
    1. NO SALUDES, NO TE DESPIDAS, NO DES EXPLICACIONES. Genera ÚNICAMENTE el contenido de la prueba adaptada.
    2. La primera línea de tu respuesta DEBE ser el título de la prueba (Ej: Evaluación Adaptada de Lenguaje).
    3. NO incluyas campos para "Nombre", "Curso" o "Fecha" (eso ya está en la plantilla).
    4. NO uses formato Markdown como asteriscos (**) o almohadillas (#). Escribe en texto plano.
    5. Para las alternativas, usa un guion (-) al inicio de cada opción.
    6. Para las preguntas de Verdadero/Falso, deja una línea de guiones bajos (______) al inicio.
    7. Si la evaluación lo permite, incluye una instrucción que diga "Dibuja..." para evaluar la comprensión de forma gráfica.
    8. Si la evaluación lo permite, incluye una instrucción que diga "Escribe..." para evaluar caligrafía o completación.
    9. Adapta el lenguaje y la complejidad según las características del {necesidad}.

    TEXTO ORIGINAL DE LA PRUEBA:
    {texto}
    """
    
    response = modelo.generate_content(prompt)
    return response.text

# --- 6. ÁREA DE CARGA Y EJECUCIÓN ---
st.markdown('<div class="item-card">', unsafe_allow_html=True)
st.markdown("### 2. Sube tu archivo y genera")
archivo = st.file_uploader("Arrastra aquí tu prueba original (.docx)", type=["docx"])

api_key_input = api_key_configurada if api_key_configurada else st.text_input("Ingresa tu API Key:", type="password")

if archivo and st.button("🚀 Crear Prueba Profesional", use_container_width=True):
    if not api_key_input:
        st.error("Por favor, ingresa tu API Key.")
    else:
        # Intentar conectar con Gemini
        try:
            genai.configure(api_key=api_key_input)
            # Buscar el modelo correcto (gemini-3-flash-preview o similar)
            modelos_disponibles = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            
            modelo_a_usar = None
            # Buscamos primero el 3 flash preview
            for m in modelos_disponibles:
                if 'gemini-3-flash-preview' in m:
                    modelo_a_usar = m
                    break
            # Si no está, buscamos el 1.5 flash
            if not modelo_a_usar:
                 for m in modelos_disponibles:
                    if 'gemini-1.5-flash' in m:
                        modelo_a_usar = m
                        break
            
            if not modelo_a_usar:
                st.error("No se encontró un modelo compatible (Flash) en tu cuenta.")
            else:
                with st.spinner(f"Diseñando material con {modelo_a_usar.split('/')[-1]}..."):
                    # 1. Leer archivo
                    texto_original = leer_docx(archivo)
                    
                    # 2. Llamar a la IA
                    texto_ia = llamar_ia(texto_original, curso_sel, asignatura_sel, necesidad_sel, api_key_input, modelo_a_usar)
                    
                    # 3. Extraer el título para pasarlo a la función creadora
                    lineas_ia = [l.strip() for l in texto_ia.split('\n') if l.strip()]
                    titulo_documento = lineas_ia[0] if lineas_ia else f"Evaluación Adaptada - {asignatura_sel}"
                    
                    # 4. Crear el Word
                    word_final = crear_docx_adaptado(texto_ia, titulo_documento)
                    
                    st.balloons()
                    st.success("✨ ¡Material adaptado con éxito!")
                    st.download_button(
                        label="⬇️ Descargar Documento Adaptado", 
                        data=word_final, 
                        file_name=f"PACI_{asignatura_sel}_{curso_sel}.docx", 
                        use_container_width=True,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    with st.expander("Ver texto generado por la IA"):
                        st.text(texto_ia) # Usamos st.text para ver el texto plano sin formatos raros

        except Exception as e:
            st.error(f"Ocurrió un error: {e}")
            st.info("Sugerencia: Verifica que tu API Key sea correcta y tenga permisos para usar los modelos Gemini.")

st.markdown('</div>', unsafe_allow_html=True)
